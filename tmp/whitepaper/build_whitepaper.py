from __future__ import annotations

import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/bryllim/Developer/egov-agent")
OUT_DOCX = ROOT / "output/docx/eGov-Agent-Technical-Whitepaper.docx"
WORK = ROOT / "tmp/whitepaper"
ARCH_IMG = WORK / "architecture.png"
FLOW_IMG = WORK / "request-lifecycle.png"
TRUST_IMG = WORK / "trust-boundaries.png"
LOGO = ROOT / "public/egov-logo-hq.png"

VERSION = "1.0"
BUILD = "31ea953"
DATE = "22 July 2026"

# Named preset override: narrative_proposal_a4_whitepaper
# Base: narrative_proposal. Overrides are deliberate for a print-first A4 technical paper.
PAGE_W = Cm(21.0)
PAGE_H = Cm(29.7)
MARGIN_X = Inches(0.75)
MARGIN_TOP = Inches(0.72)
MARGIN_BOTTOM = Inches(0.68)
CONTENT_IN = 8.2677 - 1.5
CONTENT_DXA = round(CONTENT_IN * 1440)
TABLE_INDENT_DXA = 90

FONT = "Arial"
MONO = "Courier New"

NAVY = "063D7D"
BLUE = "0A4F9E"
SKY = "2F89E6"
INK = "142033"
MUTED = "5F6B7A"
LIGHT = "F3F7FC"
PALE = "EAF2FC"
LINE = "D8E2EF"
GREEN = "087A55"
GREEN_BG = "EAF7F2"
GOLD = "9A6500"
GOLD_BG = "FFF5D9"
RED = "A02C2C"
RED_BG = "FCEEEE"
WHITE = "FFFFFF"
GRAY_BG = "F6F7F9"

FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=LINE, size=5, *, top=True, bottom=True, start=True, end=True) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, enabled in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        tag = f"w:{edge}"
        node = tc_borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_borders.append(node)
        node.set(qn("w:val"), "single" if enabled else "nil")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=TABLE_INDENT_DXA) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def set_font(run, *, size=None, bold=None, color=None, italic=None, name=FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_paragraph_border(paragraph, *, side="bottom", color=LINE, size=6, space=4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        for edge in ("top", "start", "bottom", "end"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "5")
            node.set(qn("w:space"), "5")
            node.set(qn("w:color"), border)
            p_bdr.append(node)
        p_pr.append(p_bdr)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    r_pr.append(sz)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    base_abs = max(existing_abs or [0]) + 1
    base_num = max(existing_num or [0]) + 1

    def make_abstract(abs_id: int, fmt: str, text: str, left: int, hanging: int) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "70")
        spacing.set(qn("w:line"), "270")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT)
        r_fonts.set(qn("w:hAnsi"), FONT)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def make_num(num_id: int, abs_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abs_id))
        num.append(abstract_num_id)
        numbering.append(num)

    make_abstract(base_abs, "bullet", "•", 500, 240)
    make_num(base_num, base_abs)
    make_abstract(base_abs + 1, "decimal", "%1.", 520, 260)
    make_num(base_num + 1, base_abs + 1)
    return base_num, base_num + 1


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    n = OxmlElement("w:numId")
    n.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, n])
    p_pr.append(num_pr)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5.2)
    normal.paragraph_format.line_spacing = 1.13
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in (
        ("Heading 1", 17.5, NAVY, 14, 6),
        ("Heading 2", 12.5, BLUE, 9, 4),
        ("Heading 3", 10.5, INK, 7, 3),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.font.size = Pt(7.7)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = False


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.left_margin = MARGIN_X
    section.right_margin = MARGIN_X
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("eGOV AGENT  /  TECHNICAL WHITEPAPER")
    set_font(r, size=7.2, bold=True, color=MUTED)
    set_paragraph_border(p, color=LINE, size=4, space=3)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(CONTENT_IN))
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [round(CONTENT_DXA * 0.72), CONTENT_DXA - round(CONTENT_DXA * 0.72)], indent_dxa=0)
    for cell in table.rows[0].cells:
        set_cell_borders(cell, top=False, bottom=False, start=False, end=False)
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    r = left.add_run(f"Version {VERSION}  |  {DATE}  |  Build {BUILD}")
    set_font(r, size=7.2, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    r = right.add_run("PAGE ")
    set_font(r, size=7.2, bold=True, color=MUTED)
    add_page_field(right)
    for run in right.runs:
        set_font(run, size=7.2, bold=True, color=MUTED)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"TECHNICAL EVALUATION EDITION  |  VERSION {VERSION}")
    set_font(r, size=7.4, bold=True, color=MUTED)


def add_heading(doc: Document, text: str, level=1, kicker: str | None = None) -> None:
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(kicker.upper())
        set_font(r, size=7.6, bold=True, color=SKY)
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, after=5.2, color=INK, size=9.4, align=None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=size, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=size, color=color)
    else:
        r = p.add_run(text)
        set_font(r, size=size, color=color)


def add_bullet(doc: Document, text: str, bullet_num_id: int, *, size=9.1, color=INK, bold_prefix: str | None = None, after=3.5) -> None:
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=size, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=size, color=color)
    else:
        r = p.add_run(text)
        set_font(r, size=size, color=color)


def add_number(doc: Document, text: str, number_num_id: int, *, size=9.1, after=3.5) -> None:
    p = doc.add_paragraph()
    apply_num(p, number_num_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_font(r, size=size, color=INK)


def add_callout(doc: Document, label: str, text: str, *, tone="blue", compact=False) -> None:
    fill, border, label_color = {
        "blue": (LIGHT, "BCD4F2", BLUE),
        "green": (GREEN_BG, "B9E3D3", GREEN),
        "gold": (GOLD_BG, "E7D394", GOLD),
        "red": (RED_BG, "E7B7B7", RED),
    }[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7 if not compact else 5)
    p.paragraph_format.line_spacing = 1.08
    shade_paragraph(p, fill, border)
    r = p.add_run(label.upper() + "  ")
    set_font(r, size=7.8 if compact else 8.1, bold=True, color=label_color)
    r = p.add_run(text)
    set_font(r, size=8.8 if compact else 9.2, color=INK)


def add_data_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float], *, header_fill=BLUE, font_size=8.1, status_col: int | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    widths_dxa = [round(CONTENT_DXA * w) for w in widths]
    widths_dxa[-1] += CONTENT_DXA - sum(widths_dxa)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell, color=header_fill, size=5)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_font(r, size=8.0, bold=True, color=WHITE)
    for r_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cell = cells[i]
            set_cell_shading(cell, WHITE if r_idx % 2 == 0 else "F8FAFD")
            set_cell_borders(cell, color=LINE, size=4)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            color = INK
            if status_col == i:
                lower = text.lower()
                color = GREEN if any(k in lower for k in ("implemented", "ready", "required", "go")) else GOLD if any(k in lower for k in ("simulated", "proposed", "pilot")) else INK
            r = p.add_run(text)
            set_font(r, size=font_size, bold=(status_col == i), color=color)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, text: str, *, size=7.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(7)
    p.paragraph_format.right_indent = Pt(7)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.0
    shade_paragraph(p, GRAY_BG, LINE)
    r = p.add_run(text)
    set_font(r, size=size, color=INK, name=MONO)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_picture(doc: Document, path: Path, width_in: float, alt: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width_in))
    inline._inline.docPr.set("descr", alt)
    inline._inline.docPr.set("title", alt)


def new_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def page_label(doc: Document, n: str, label: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{n}  /  {label.upper()}")
    set_font(r, size=7.3, bold=True, color=SKY)


def load_font(size: int, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size=size)


def rounded(draw, box, radius, fill, outline=None, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text_box(draw, box, title, lines, *, fill="#ffffff", outline="#cddcec", title_color="#063d7d", font_scale=1.0, align="left"):
    x1, y1, x2, y2 = box
    rounded(draw, box, 22, fill, outline, 3)
    # Enforce a print-safe floor: these diagrams are placed at ~6.7 inches wide.
    t_font = load_font(max(30, round(38 * font_scale)), True)
    b_font = load_font(max(24, round(30 * font_scale)), False)
    pad = 28
    if align == "center":
        tb = draw.textbbox((0, 0), title, font=t_font)
        tx = x1 + (x2 - x1 - (tb[2] - tb[0])) / 2
    else:
        tx = x1 + pad
    draw.text((tx, y1 + 22), title, font=t_font, fill=title_color)
    y = y1 + 75
    for line in lines:
        if align == "center":
            bb = draw.textbbox((0, 0), line, font=b_font)
            lx = x1 + (x2 - x1 - (bb[2] - bb[0])) / 2
        else:
            lx = x1 + pad
        draw.text((lx, y), line, font=b_font, fill="#465568")
        y += max(34, round(40 * font_scale))


def arrow(draw, start, end, color="#0a4f9e", width=7):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 22
    for delta in (2.55, -2.55):
        p = (end[0] + length * math.cos(angle + delta), end[1] + length * math.sin(angle + delta))
        draw.line([end, p], fill=color, width=width)


def build_architecture() -> None:
    img = Image.new("RGB", (2400, 1480), "white")
    d = ImageDraw.Draw(img)
    title = load_font(42, True)
    band = load_font(24, True)
    d.text((70, 38), "Production reference architecture", font=title, fill="#142033")
    bands = [
        ("CHANNELS", 135, 355, "#f5f8fc"),
        ("ORCHESTRATION CONTROL PLANE", 390, 760, "#edf4fd"),
        ("TRUSTED DATA SERVICES", 795, 1090, "#f4f8f7"),
        ("GOVERNMENT SYSTEMS OF RECORD", 1125, 1405, "#fff8e8"),
    ]
    for label, y1, y2, fill in bands:
        rounded(d, (50, y1, 2350, y2), 28, fill, "#d7e1ee", 3)
        d.text((76, y1 + 18), label, font=band, fill="#5f6b7a")

    # Channels
    for i, (title_t, sub) in enumerate([
        ("eGovPH", ["Web / mobile"]),
        ("Messaging", ["Viber / future channels"]),
        ("Assisted service", ["Kiosk / call center"]),
    ]):
        x = 150 + i * 730
        text_box(d, (x, 200, x + 560, 325), title_t, sub, fill="#ffffff")

    # Control plane
    control = [
        ("API Gateway", ["OIDC session", "rate limits"], 105, 485),
        ("Agent Orchestrator", ["plan + tool calls", "workflow state"], 580, 485),
        ("Policy & Consent", ["risk tier", "approval gates"], 1055, 485),
        ("Adapter Registry", ["versioned contracts", "agency capabilities"], 1530, 485),
    ]
    for title_t, lines, x, y in control:
        text_box(d, (x, y, x + 375, y + 180), title_t, lines, fill="#ffffff", font_scale=.82)
    for i in range(3):
        arrow(d, (480 + i * 475, 575), (565 + i * 475, 575))
    text_box(d, (650, 680, 1740, 755), "Durable workflow + event bus", [], fill="#dcecff", font_scale=.72, align="center")

    # Data services
    for i, (title_t, lines) in enumerate([
        ("Memory", ["provenance + TTL"]),
        ("Vault", ["encrypted objects"]),
        ("Audit", ["append-only events"]),
        ("Observability", ["logs / metrics / traces"]),
    ]):
        x = 105 + i * 560
        text_box(d, (x, 885, x + 435, 1040), title_t, lines, fill="#ffffff", outline="#cce4dc", font_scale=.82)

    # Government systems
    gov = [
        ("National ID eVerify", ["identity authentication"]),
        ("eGovDX / SSO", ["platform interoperability"]),
        ("eGov Pay + eReceipt", ["payment and proof"]),
        ("Agency / LGU APIs", ["systems of record"]),
    ]
    for i, (title_t, lines) in enumerate(gov):
        x = 105 + i * 560
        text_box(d, (x, 1210, x + 435, 1365), title_t, lines, fill="#ffffff", outline="#e5d3a2", font_scale=.75)

    # vertical arrows
    for x in (430, 890, 1370, 1840):
        arrow(d, (x, 1080), (x, 1195), color="#9a6500", width=6)
    img.save(ARCH_IMG, quality=95)


def build_flow() -> None:
    img = Image.new("RGB", (2400, 1260), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 38), "High-risk transaction lifecycle", font=load_font(42, True), fill="#142033")
    steps = [
        ("1", "Authenticate", "Short-lived\nidentity assertion"),
        ("2", "Understand", "Intent + minimum\nrequired facts"),
        ("3", "Plan", "Allowed tools +\nrisk classification"),
        ("4", "Approve", "Explicit consent\nfor side effects"),
        ("5", "Execute", "Idempotent adapter\ncall"),
        ("6", "Verify", "Read back from\nsystem of record"),
        ("7", "Issue proof", "Agency receipt /\nacknowledgement"),
        ("8", "Reconcile", "Audit event +\nasync recovery"),
    ]
    positions = []
    for i in range(4):
        positions.append((100 + i * 560, 170))
    for i in range(4):
        positions.append((1780 - i * 560, 720))
    for idx, ((num, title_t, body), (x, y)) in enumerate(zip(steps, positions)):
        fill = "#eaf2fc" if idx < 4 else "#eaf7f2"
        outline = "#b8d1ef" if idx < 4 else "#b9e3d3"
        rounded(d, (x, y, x + 470, y + 300), 28, fill, outline, 3)
        rounded(d, (x + 28, y + 28, x + 96, y + 96), 34, "#0a4f9e" if idx < 4 else "#087a55")
        bb = d.textbbox((0, 0), num, font=load_font(30, True))
        d.text((x + 62 - (bb[2]-bb[0])/2, y + 62 - (bb[3]-bb[1])/2 - 2), num, font=load_font(30, True), fill="white")
        d.text((x + 28, y + 125), title_t, font=load_font(38, True), fill="#142033")
        for li, line in enumerate(body.split("\n")):
            d.text((x + 28, y + 190 + li * 39), line, font=load_font(30), fill="#4c5b6d")
    for i in range(3):
        arrow(d, (positions[i][0] + 470, positions[i][1] + 150), (positions[i+1][0] - 15, positions[i+1][1] + 150))
    arrow(d, (positions[3][0] + 235, positions[3][1] + 300), (positions[4][0] + 235, positions[4][1] - 15))
    for i in range(4, 7):
        arrow(d, (positions[i][0], positions[i][1] + 150), (positions[i+1][0] + 485, positions[i+1][1] + 150))
    rounded(d, (260, 1090, 2140, 1200), 24, "#fff5d9", "#e7d394", 3)
    msg = "No UI success state is authoritative: every consequential action is read back from the agency ledger before completion is shown."
    bb = d.textbbox((0, 0), msg, font=load_font(25, True))
    d.text((1200 - (bb[2]-bb[0])/2, 1128), msg, font=load_font(25, True), fill="#7b5200")
    img.save(FLOW_IMG, quality=95)


def build_trust() -> None:
    img = Image.new("RGB", (2400, 1170), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 38), "Trust boundaries and data minimization", font=load_font(42, True), fill="#142033")
    zones = [
        (80, 150, 690, 1050, "Citizen device", "#f5f8fc", "#cddcec"),
        (740, 150, 1660, 1050, "Government control plane", "#edf4fd", "#a9c9ee"),
        (1710, 150, 2320, 1050, "External AI provider (optional)", "#f8f8f8", "#c7c7c7"),
    ]
    for x1, y1, x2, y2, title_t, fill, outline in zones:
        rounded(d, (x1, y1, x2, y2), 30, fill, outline, 4)
        d.text((x1 + 30, y1 + 28), title_t, font=load_font(31, True), fill="#142033")
    for i, (title_t, lines) in enumerate([
        ("Session", ["short-lived token", "no agency secrets"]),
        ("Consent UI", ["purpose + scope", "revocable choices"]),
        ("Local cache", ["minimal", "time-bounded"]),
    ]):
        text_box(d, (130, 260 + i*235, 640, 440 + i*235), title_t, lines, fill="white", font_scale=.78)
    for i, (title_t, lines) in enumerate([
        ("Identity broker", ["OIDC / eVerify assertion"]),
        ("Policy engine", ["purpose / role / risk"]),
        ("Encrypted vault", ["KMS keys + malware scan"]),
        ("Audit ledger", ["append-only, correlated"]),
        ("Agency adapters", ["delegated credentials"]),
    ]):
        x = 790 + (i % 2) * 420
        y = 260 + (i // 2) * 235
        text_box(d, (x, y, x + 370, y + 180), title_t, lines, fill="white", outline="#bdd4f0", font_scale=.68)
    text_box(d, (1760, 300, 2270, 500), "Redacted prompt", ["no raw IDs", "minimum context"], fill="white", outline="#c7c7c7", font_scale=.75)
    text_box(d, (1760, 620, 2270, 820), "Model response", ["untrusted input", "schema-validated"], fill="white", outline="#c7c7c7", font_scale=.75)
    arrow(d, (1655, 400), (1745, 400), color="#6b7280", width=6)
    arrow(d, (1745, 740), (1655, 740), color="#6b7280", width=6)
    d.text((1770, 900), "Optional path only for complex tasks", font=load_font(22, True), fill="#6b7280")
    d.text((1770, 935), "under approved zero-retention terms", font=load_font(22), fill="#6b7280")
    img.save(TRUST_IMG, quality=95)


def build_document() -> None:
    build_architecture()
    build_flow()
    build_trust()

    doc = Document()
    setup_styles(doc)
    configure_page(doc)
    bullet_id, number_id = configure_numbering(doc)
    doc.core_properties.title = "eGov Agent: A Feasible Agentic Orchestration Layer for eGovPH"
    doc.core_properties.subject = "Technical whitepaper for production feasibility review"
    doc.core_properties.author = "eGov Agent Team"
    doc.core_properties.keywords = "eGovPH, agentic AI, government services, interoperability, eGovPay, eVerify"
    doc.core_properties.comments = "Generated from audited prototype build 31ea953."

    # 01 Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    inline = run.add_picture(str(LOGO), width=Inches(0.68))
    inline._inline.docPr.set("descr", "eGovPH ring mark")
    inline._inline.docPr.set("title", "eGovPH")
    wr = p.add_run("  eGov ")
    set_font(wr, size=17, bold=True, color=INK)
    wr = p.add_run("Agent")
    set_font(wr, size=17, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("TECHNICAL WHITEPAPER")
    set_font(r, size=8.5, bold=True, color=SKY)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("A Feasible Agentic\nOrchestration Layer\nfor eGovPH")
    set_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Using the government connections that already exist - and completing the work for the citizen.")
    set_font(r, size=13.5, color=MUTED)

    add_callout(
        doc,
        "Thesis",
        "eGovPH already provides the digital front door and shared rails. eGov Agent adds a consent-aware execution layer that can understand a request, coordinate approved services, and return one verifiable outcome without moving agency ownership of records.",
        tone="blue",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"VERSION {VERSION}  /  {DATE}")
    set_font(r, size=8.3, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Prototype build {BUILD}  /  Technical evaluation edition")
    set_font(r, size=8.3, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_border(p, color=SKY, size=16, space=0)

    # 02 Executive summary
    new_page(doc)
    page_label(doc, "01", "Executive summary")
    add_heading(doc, "Executive summary", 1)
    add_body(doc, "eGov Agent is a proposed orchestration layer for eGovPH: a citizen asks for an outcome in plain language, the system identifies the responsible services, asks for narrowly scoped consent, executes through approved government APIs, and returns a synchronized status, document, appointment, report acknowledgement, or receipt.")
    add_body(doc, "The proposition is deliberately incremental. It reuses existing platforms - including eGovDX, National ID eVerify, eGov Pay, eReceipt, eReport, and agency or LGU systems of record - rather than building a parallel identity, payment, or records platform. DICT already presents these components as part of the national e-government ecosystem.[1]")
    add_heading(doc, "What changes for the citizen", 2)
    for text in [
        "One request replaces manual discovery across agency portals.",
        "Verified profile facts and previously approved documents can be reused with provenance instead of re-entered.",
        "Cross-agency work is represented as one auditable workflow, not a chain of opaque redirects.",
        "Payments are completed through eGov Pay and are shown as final only after agency-side verification.",
        "High-risk actions remain consented, policy-checked, reversible where possible, and attributable.",
    ]:
        add_bullet(doc, text, bullet_id)

    add_heading(doc, "Recommendation", 2)
    add_callout(doc, "Proceed with a bounded pilot", "Start with one read-only service, one appointment or filing service, one eGovPay transaction, and one LGU eReport workflow. Require sandbox APIs, a Privacy Impact Assessment, contract tests, explicit consent receipts, and measurable service-level objectives before adding more agencies.", tone="green")
    add_heading(doc, "Evidence standard", 2)
    add_body(doc, "This paper separates implemented user experience from simulated integrations and proposed production components. The current repository is a strong interaction prototype; it is not a connected government production system. No live agency credentials, resident database, payment settlement, or emergency dispatch is claimed.")

    # 03 Prototype boundary
    new_page(doc)
    page_label(doc, "02", "Prototype boundary")
    add_heading(doc, "What exists today - and what does not", 1)
    add_body(doc, "The audited build is a Next.js 16.2.10 and React 19.2.4 client application. Its interaction model is implemented; its government outcomes are deterministic simulations. This is appropriate for a hackathon proof of experience, provided the production gap is explicit.")
    add_data_table(
        doc,
        ["Capability", "Current build", "Production requirement"],
        [
            ["Conversational service discovery", "Implemented", "Replace scripted intent matching with policy-governed routing and evaluated models."],
            ["Agent activity trace", "Implemented", "Populate from real distributed traces and adapter events."],
            ["File stamps, previews, uploads", "Implemented", "Object storage, malware scanning, encryption, retention, DLP."],
            ["Personal memory", "Simulated", "Provenance-first store, consent, expiry, correction, deletion, source reconciliation."],
            ["National ID authentication", "Simulated", "PSA relying-party onboarding and NIDAS/eVerify integration."],
            ["Agency records and appointments", "Simulated", "Agency-approved APIs, delegated authorization, sandbox and UAT."],
            ["eGovPay and eReceipt", "Simulated", "Server-side order creation, signed callback/polling, ledger reconciliation."],
            ["eReport dispatch", "Simulated", "Authoritative routing directory, agency queues, acknowledgements, escalation rules."],
            ["Backend, database, queue", "Not present", "Government-hosted control plane and durable workflow engine."],
        ],
        [0.24, 0.17, 0.59],
        status_col=1,
        font_size=7.8,
    )
    add_heading(doc, "Auditable implementation footprint", 2)
    add_data_table(
        doc,
        ["Source", "Purpose"],
        [
            ["app/agent/brain.ts", "Typed cards, scripted intent-to-plan mapping, trace steps, vault references."],
            ["app/agent/page.tsx", "Conversation UI, attachments, voice entry, agent execution animation."],
            ["app/agent/forms.ts", "Printable forms, appointment passes, simulated receipts and eReport acknowledgement."],
            ["app/agent/memory/page.tsx", "Source-labelled personal context interface."],
            ["app/agent/vault/page.tsx", "Document vault interactions, file previews and uploads."],
        ],
        [0.31, 0.69],
        header_fill=NAVY,
        font_size=8.0,
    )
    add_callout(doc, "Critical boundary", "The current UI must never be connected directly to agency APIs. All tokens, signatures, policy checks, payment verification, document issuance, and audit events belong in a server-side government control plane.", tone="red", compact=True)

    # 04 Architecture
    new_page(doc)
    page_label(doc, "03", "Reference architecture")
    add_heading(doc, "A thin orchestration layer over existing rails", 1)
    add_body(doc, "The target design keeps agency systems authoritative. eGov Agent stores only workflow state, user-approved context, and cryptographic or audit evidence required to complete and explain a transaction. It does not replicate agency registries into a general-purpose AI database.")
    add_picture(doc, ARCH_IMG, 6.68, "Production reference architecture showing channels, orchestration control plane, trusted data services, and government systems of record")
    add_caption(doc, "Figure 1. Production reference architecture. Arrows represent authenticated, policy-governed API calls; agency systems remain the systems of record.")
    add_heading(doc, "Core design decisions", 2)
    for text in [
        "Control plane, not monolith: the orchestrator coordinates identity, policy, adapters, and workflow state; it does not absorb agency business logic.",
        "Adapter isolation: every agency integration is versioned, rate-limited, observable, and removable without changing the conversation layer.",
        "Durable execution: long-running work uses a workflow engine and event bus so retries do not duplicate bookings, payments, or reports.",
        "Minimal AI privilege: models propose structured plans; a deterministic policy engine decides which tools and data scopes are allowed.",
    ]:
        add_bullet(doc, text, bullet_id, size=8.85)

    # 05 Lifecycle
    new_page(doc)
    page_label(doc, "04", "Transaction lifecycle")
    add_heading(doc, "From request to verifiable outcome", 1)
    add_body(doc, "Government transactions are distributed workflows. Network timeouts, partial success, manual review, and delayed agency processing are normal conditions. The design therefore treats a conversation as a durable state machine rather than one model response.")
    add_picture(doc, FLOW_IMG, 6.66, "Eight-step high-risk transaction lifecycle from authentication through reconciliation")
    add_caption(doc, "Figure 2. Every consequential action passes through authentication, policy, explicit approval, idempotent execution, authoritative read-back, and reconciliation.")
    add_heading(doc, "Execution invariants", 2)
    add_data_table(
        doc,
        ["Invariant", "Engineering rule"],
        [
            ["No duplicate side effects", "Client supplies an idempotency key; adapter persists result before acknowledging."],
            ["No model-issued authority", "Only allow-listed tools can cause effects; policy evaluation is deterministic."],
            ["No inferred payment success", "Success requires a verified gateway event plus agency ledger confirmation."],
            ["No silent scope expansion", "A new agency, purpose, data category, or fee requires renewed approval."],
            ["No lost partial state", "Workflow checkpoints and compensating actions survive process restarts."],
        ],
        [0.28, 0.72],
        font_size=8.15,
    )
    add_callout(doc, "Delivery semantics", "Distributed systems cannot promise exactly-once delivery across independent agencies. The feasible contract is at-least-once messaging with idempotent agency operations, deterministic deduplication, and reconciliation.", tone="gold", compact=True)

    # 06 Contracts
    new_page(doc)
    page_label(doc, "05", "Integration contracts")
    add_heading(doc, "A stable adapter contract for uneven agency systems", 1)
    add_body(doc, "Agencies differ in authentication, data models, turnaround times, and legacy constraints. eGov Agent normalizes the orchestration contract while preserving the agency-specific implementation behind each adapter.")
    add_code_block(doc, """POST /v1/actions/execute\nIdempotency-Key: 01J...\nAuthorization: Bearer <delegated-token>\n\n{\n  \"action\": \"appointment.reserve\",\n  \"agency\": \"DFA\",\n  \"subject_ref\": \"opaque-user-ref\",\n  \"consent_receipt\": \"cr_...\",\n  \"input\": {\"site_id\": \"...\", \"slot_id\": \"...\"}\n}\n\n202 Accepted\n{\"operation_id\": \"op_...\", \"status\": \"PENDING\", \"poll_after_s\": 5}""", size=7.25)
    add_data_table(
        doc,
        ["Concern", "Contract"],
        [
            ["Authentication", "OIDC/OAuth 2.1 where available; short-lived delegated token; mTLS between services."],
            ["Idempotency", "Required for every side effect; scope key to subject + agency + action + request hash."],
            ["Async operations", "202 + operation ID; webhook or polling; terminal states are explicit."],
            ["Errors", "Stable codes: validation, consent, auth, rate limit, unavailable, manual review, rejected."],
            ["Versioning", "Semantic adapter versions; deprecation window; contract tests in agency sandbox."],
            ["Evidence", "Correlation ID, policy decision ID, consent receipt, agency reference, timestamps."],
        ],
        [0.24, 0.76],
        font_size=8.0,
    )
    add_heading(doc, "Real onboarding path", 2)
    add_body(doc, "DICT's published Citizen's Charter describes eGovPH integration as a government-to-government process involving a Letter of Intent, eGovPH Request Form, Memorandum of Understanding, technical integration, internal testing, UAT, and deployment; timing depends on client readiness.[8] The implementation plan in this paper follows that path rather than assuming public, self-service APIs.")
    add_callout(doc, "Feasibility principle", "Pilot only with agencies that can provide an accountable product owner, sandbox or test double, canonical service rules, and a named incident/escalation path.", tone="green", compact=True)

    # 07 Identity/memory/vault
    new_page(doc)
    page_label(doc, "06", "Identity, memory, and vault")
    add_heading(doc, "Personalization without a surveillance graph", 1)
    add_body(doc, "National ID eVerify is a realistic identity rail, but integration requires relying-party onboarding. PSA describes eVerify as an authentication service for government and private relying parties and notes regulatory and technical onboarding, including API access.[3][4] eGov Agent should consume a scoped identity assertion, not store biometrics or the PhilSys registry.")
    add_heading(doc, "Recommended memory architecture", 2)
    add_data_table(
        doc,
        ["Store", "Use", "Why"],
        [
            ["Relational profile facts", "Address city, language, notification preference", "Queryable, correctable, source-linked, easy to expire."],
            ["Limited relationship graph", "Household, authorized representative, document-to-case links", "Useful where relationships matter; avoid a universal knowledge graph."],
            ["Vector index", "Semantic retrieval over user-approved documents", "Find relevant text; never the authoritative fact store."],
            ["Agency references", "Opaque IDs and last-known state", "Point back to systems of record rather than copying registries."],
        ],
        [0.25, 0.37, 0.38],
        font_size=7.85,
    )
    add_body(doc, "Every remembered fact should carry: subject, value, source, purpose, confidence, captured_at, expires_at, consent scope, sensitivity, and supersedes. Facts synchronized from agencies are not editable as truth; the citizen can flag them for correction at the source.")
    add_heading(doc, "Vault controls", 2)
    for text in [
        "Encrypted object storage with envelope encryption and per-tenant keys managed by a government KMS.",
        "Malware scanning, MIME validation, content-disarm policy, and quarantine before any model or adapter sees a file.",
        "Signed, short-lived download URLs; watermarking for generated previews; no public object paths.",
        "Retention by purpose and document type, with legal hold and user deletion workflows where allowed.",
        "Document provenance: issuer, hash, verification status, valid-from/to, and each transaction that reused it.",
    ]:
        add_bullet(doc, text, bullet_id, size=8.9)
    add_callout(doc, "Answer to the knowledge-graph question", "Use a graph selectively for real relationships. A graph should not become the primary store for all personal context; a hybrid model is easier to govern, correct, delete, and reconcile with authoritative agencies.", tone="blue", compact=True)

    # 08 Payments
    new_page(doc)
    page_label(doc, "07", "Payments and eReceipt")
    add_heading(doc, "eGovPay integration that can survive reconciliation", 1)
    add_body(doc, "The public eGovPay site describes both a Payment Gateway API and Payment Links, with cards, online bank transfers, and e-wallets.[2] The production design should use a hosted or government-controlled checkout so the agent never handles payment credentials.")
    add_heading(doc, "Authoritative payment flow", 2)
    for text in [
        "The agency creates an order server-side with service code, assessed amount, expiry, citizen reference, and idempotency key.",
        "eGovPay returns a hosted checkout reference. The citizen reviews the exact agency, service, amount, and refund rules before authorization.",
        "A signed webhook and/or server-to-server status query reports gateway settlement. Redirect success is never sufficient.",
        "The adapter validates merchant, order, amount, currency, state transition, timestamp, and replay nonce.",
        "The agency posts the payment to its ledger and returns the official receipt or eReceipt reference.",
        "Only then does eGov Agent mark the transaction paid and display the immutable receipt copy.",
    ]:
        add_number(doc, text, number_id, size=8.8, after=3.0)
    add_data_table(
        doc,
        ["State", "Meaning", "Citizen-facing behavior"],
        [
            ["CREATED", "Order exists; no payment", "Show fee and expiry."],
            ["PENDING", "Authorization or settlement in progress", "Do not resubmit; allow status refresh."],
            ["PAID_UNPOSTED", "Gateway paid; agency ledger pending", "Show processing; start reconciliation."],
            ["PAID", "Agency ledger confirmed", "Release official proof and continue service."],
            ["FAILED / EXPIRED", "No settlement", "Allow a new order with a new idempotency key."],
            ["REFUND_PENDING / REFUNDED", "Reversal workflow", "Show agency-owned refund reference."],
        ],
        [0.22, 0.34, 0.44],
        font_size=7.75,
    )
    add_callout(doc, "Non-negotiable", "The agent may format and present a receipt, but it must not invent an official receipt number. The issuing agency or the designated eReceipt service must return that identifier after verified posting.", tone="red", compact=True)

    # 09 eReport
    new_page(doc)
    page_label(doc, "08", "eReport")
    add_heading(doc, "A credible eReport workflow: assist, confirm, dispatch, track", 1)
    add_body(doc, "The prototype's flooding scenario demonstrates the intended experience: accept a photo and description, infer likely incident type, resolve location, identify responsible desks, ask the citizen to review, dispatch once, and synchronize acknowledgements. Production safety requires the AI to assist triage - never to impersonate a dispatcher or claim acknowledgement without an agency event.")
    add_data_table(
        doc,
        ["Stage", "Automation", "Required guardrail"],
        [
            ["Intake", "Extract description, time, media metadata", "Disclose metadata use; strip unnecessary EXIF on onward sharing."],
            ["Location", "Geocode and jurisdiction match", "Citizen confirms pin; retain confidence and resolver version."],
            ["Triage", "Suggest category and severity", "Human/agency rules override model; no definitive emergency diagnosis."],
            ["Routing", "Match primary and secondary responders", "Authoritative service directory; deduplicate by time/place/evidence."],
            ["Dispatch", "Create agency work items", "Idempotent fan-out; per-agency acceptance or rejection captured."],
            ["Tracking", "Aggregate status and ETA", "Only display agency-supplied status; timestamp every update."],
            ["Closure", "Ask for outcome feedback", "Agency closes official case; citizen may dispute or reopen."],
        ],
        [0.16, 0.38, 0.46],
        font_size=7.65,
    )
    add_heading(doc, "Report state model", 2)
    add_code_block(doc, "DRAFT -> CONFIRMED -> DISPATCHING -> DISPATCHED_PARTIAL|DISPATCHED\n      -> ACKNOWLEDGED -> IN_PROGRESS -> RESOLVED -> CLOSED\nAny state -> REJECTED|DUPLICATE|CANCELLED; every transition records actor, reason, and timestamp.", size=7.5)
    add_heading(doc, "Public-safety safeguards", 2)
    for text in [
        "Surface the official emergency channel when immediate danger is indicated; the form must not delay emergency contact.",
        "Rate-limit and abuse-detect without silencing legitimate repeated reports during disasters.",
        "Blur faces and plates in previews by default; provide unredacted evidence only to authorized responders when lawful and necessary.",
        "Design for degraded connectivity: resumable uploads, SMS acknowledgement, offline draft, and eventual dispatch.",
    ]:
        add_bullet(doc, text, bullet_id, size=8.75)
    add_callout(doc, "Prototype truth", "The Mandaluyong CDRRMO, barangay, MMDA, and DPWH statuses shown in the current experience are simulated. Production must render each acknowledgement from an authenticated agency callback.", tone="gold", compact=True)

    # 10 Security
    new_page(doc)
    page_label(doc, "09", "Security and privacy")
    add_heading(doc, "Security model for a high-trust government agent", 1)
    add_body(doc, "The Data Privacy Act requires reasonable organizational, physical, and technical safeguards.[5] NPC Circular 2023-06 is the current cross-sector security circular and adds privacy-by-design/default and business-continuity expectations; it superseded the older government-only Circular 16-01.[6] A PIA, data-sharing agreements, records of processing, incident response, and agency DPO involvement are preconditions, not post-launch paperwork.")
    add_picture(doc, TRUST_IMG, 6.66, "Trust boundary diagram separating citizen device, government control plane, and optional external AI provider")
    add_caption(doc, "Figure 3. Raw identity, documents, credentials, and audit evidence remain inside the government control plane. Optional external models receive only approved minimum context.")
    add_data_table(
        doc,
        ["Threat", "Primary controls"],
        [
            ["Account takeover", "Phishing-resistant MFA, device/session risk, step-up authentication for high-risk actions."],
            ["Prompt injection in files", "Treat content as untrusted, isolate parsing, tool allow-list, schema validation, no instruction inheritance."],
            ["Excessive agency access", "ABAC by purpose, data category, role, consent receipt, and transaction state."],
            ["Credential leakage", "Server-only secrets, KMS/HSM, rotation, mTLS, workload identity, no secrets in prompts or logs."],
            ["Duplicate or replayed effects", "Idempotency keys, signed nonces, monotonic state machine, reconciliation."],
            ["Insider misuse", "Dual control for privileged actions, immutable audit, anomaly detection, access reviews."],
            ["Model data leakage", "Minimization, redaction, approved tenancy, zero-retention terms, output DLP."],
        ],
        [0.24, 0.76],
        font_size=7.65,
    )

    # 11 Reliability
    new_page(doc)
    page_label(doc, "10", "Reliability and operations")
    add_heading(doc, "Designed for partial failure, not perfect dependencies", 1)
    add_body(doc, "The control plane should remain useful when one agency is unavailable. It must preserve the citizen's request, show which step is delayed, avoid duplicate effects, and resume from a durable checkpoint when the dependency recovers.")
    add_data_table(
        doc,
        ["Layer", "Pattern", "Pilot target"],
        [
            ["API edge", "Stateless replicas, rate limits, WAF, request correlation", "99.9% monthly availability"],
            ["Workflow", "Durable state, queue, retries with jitter, dead-letter queue", "No lost accepted jobs"],
            ["Adapters", "Timeouts, circuit breakers, bulkheads, per-agency quotas", "Failure isolated by adapter"],
            ["Payments", "Outbox/inbox, ledger read-back, daily reconciliation", "Zero unreconciled duplicates"],
            ["Data", "Multi-AZ database, point-in-time recovery, encrypted backups", "RPO <= 15 min; RTO <= 2 h"],
            ["Observability", "OpenTelemetry traces, redacted logs, SLO alerts", ">= 95% transactions trace-complete"],
        ],
        [0.18, 0.53, 0.29],
        font_size=7.8,
    )
    add_heading(doc, "Performance budget", 2)
    for text in [
        "Conversation acknowledgement: p95 under 800 ms before long-running work begins.",
        "Policy evaluation and adapter selection: p95 under 250 ms inside the control plane.",
        "Agency actions: reported independently with their own SLA; never hidden behind a typing animation.",
        "Streaming status: server-sent events or push notifications backed by persisted workflow events.",
    ]:
        add_bullet(doc, text, bullet_id, size=8.9)
    add_heading(doc, "Capacity model", 2)
    add_code_block(doc, "Peak requests/s = monthly sessions x turns/session x peak factor / active seconds\nQueue concurrency = arrival rate x p95 dependency latency / target utilization\nStorage = workflow events + audit evidence + retained documents; scale each independently.", size=7.6)
    add_callout(doc, "Operational ownership", "Each adapter needs a runbook, dashboard, on-call owner, status page, data owner, recovery procedure, and contact at the source agency. Integration count is not a success metric if operational ownership is missing.", tone="blue", compact=True)

    # 12 Implementation
    new_page(doc)
    page_label(doc, "11", "Implementation and scalability")
    add_heading(doc, "A gated path from prototype to public pilot", 1)
    add_data_table(
        doc,
        ["Phase", "Duration", "Deliverables", "Exit gate"],
        [
            ["0 - Mobilize", "0-6 weeks", "Agency owners; PIA; data inventory; service blueprint; integration request; threat model", "Approved scope, legal basis, sandbox plan"],
            ["1 - Read-only", "6-12 weeks", "eGovPH SSO/eVerify assertion; one record lookup; audit; support console", ">=95% correct routing; no critical findings"],
            ["2 - Transact", "12-20 weeks", "Consent receipts; appointment/filing; eGovPay; reconciliation; notifications", "Zero duplicate effects; payment match >=99.95%"],
            ["3 - eReport", "16-24 weeks", "LGU routing; media handling; agency queue; acknowledgement; escalation", "Responder UAT; safety drill passed"],
            ["4 - Scale", "6-12 months", "Adapter SDK; onboarding playbook; DR; multilingual and assisted channels", "SLOs met for 90 days; independent review"],
        ],
        [0.14, 0.14, 0.43, 0.29],
        font_size=7.5,
    )
    add_heading(doc, "Recommended production stack", 2)
    add_data_table(
        doc,
        ["Concern", "Technology-neutral choice"],
        [
            ["Web and API", "Next.js front end; typed backend service; OpenAPI 3.1; government API gateway."],
            ["Workflow", "Durable workflow engine plus queue/event bus; transactional outbox."],
            ["State", "PostgreSQL-compatible relational store; object storage; Redis-compatible ephemeral cache."],
            ["Identity", "OIDC/OAuth 2.1, eGovPH SSO, eVerify relying-party integration, workload identity."],
            ["Security", "KMS/HSM, secrets manager, WAF, SIEM, SAST/DAST/SCA, container signing and policy."],
            ["AI", "Model gateway, approved local model for common tasks, optional redacted external model, evaluation harness."],
        ],
        [0.22, 0.78],
        font_size=7.75,
    )
    add_heading(doc, "AI deployment strategy", 2)
    add_body(doc, "Use deterministic rules and templates for common government intents; retrieval for service requirements; compact local models for classification and extraction; larger models only for complex language tasks. Every model output is schema-validated, scored, and subject to policy. This reduces cost and blast radius while retaining a natural interface.")
    add_callout(doc, "Accessibility and inclusion", "Target WCAG 2.2 AA, Filipino and English, screen-reader semantics, large-text support, low-bandwidth modes, human handoff, assisted-service channels, and printable outcomes from the first pilot.", tone="green", compact=True)

    # 13 Cost/benefit
    new_page(doc)
    page_label(doc, "12", "Cost and value")
    add_heading(doc, "A transparent pilot estimate - not a procurement quote", 1)
    add_body(doc, "The following range is an engineering planning estimate for a 9-12 month pilot with two national agencies, one LGU, up to 50,000 monthly active users, one payment flow, and one eReport flow. It excludes agency legacy-system remediation, statutory fees, and nationwide contact-center operations.")
    add_data_table(
        doc,
        ["Cost area", "Planning range (PHP)", "Assumption"],
        [
            ["Product and integration team", "10.5M-18.0M", "8-10 FTE equivalent; product, backend, frontend, QA, DevSecOps, design."],
            ["Security, privacy, legal, assurance", "2.0M-4.0M", "PIA, threat model, pen test, DSA/MOU support, independent review."],
            ["Cloud, observability, messaging", "1.5M-3.5M", "Pilot traffic, multi-AZ services, logs/traces, SMS/email; excludes subsidies."],
            ["Agency onboarding and change", "1.5M-3.0M", "Service mapping, sandbox/UAT, training, support runbooks."],
            ["Contingency", "2.5M-5.0M", "Legacy variance, compliance changes, recovery work."],
            ["Total pilot envelope", "18.0M-33.5M", "Range should be replaced by discovery and government procurement estimates."],
        ],
        [0.31, 0.23, 0.46],
        font_size=7.65,
    )
    add_heading(doc, "Illustrative value model", 2)
    add_data_table(
        doc,
        ["Assumption", "Illustrative input", "Derived monthly value"],
        [
            ["Citizen time avoided", "15,000 completed transactions x 45 min", "11,250 citizen-hours"],
            ["Agency handling reduced", "15,000 x 8 min", "2,000 staff-hours"],
            ["Repeat data entry avoided", "30,000 forms x 5 min", "2,500 citizen-hours"],
            ["Unit run cost", "Annual operating cost / completed transactions", "Track by service and channel"],
        ],
        [0.34, 0.36, 0.30],
        font_size=7.8,
    )
    add_body(doc, "These are hypotheses to test, not impact claims. The pilot must measure baseline journey time, completion, office visits avoided, staff touch time, payment reconciliation effort, and accessibility outcomes before declaring benefit.")
    add_heading(doc, "Cost controls", 2)
    for text in [
        "Route common intents deterministically; reserve expensive models for the small complex tail.",
        "Cache public service metadata, never sensitive personal records.",
        "Charge showback by agency adapter, model tier, notification channel, and completed transaction.",
        "Stop integrations that cannot meet operational ownership or citizen-value thresholds.",
    ]:
        add_bullet(doc, text, bullet_id, size=8.8)

    # 14 Governance/risks
    new_page(doc)
    page_label(doc, "13", "Governance and risk")
    add_heading(doc, "Risks that can make a polished agent unsafe", 1)
    add_data_table(
        doc,
        ["Risk", "Why it matters", "Mitigation / owner"],
        [
            ["Hallucinated eligibility or status", "A confident error can deny service or create false assurance.", "Rules sourced from agency; read-back; citations; agency product owner."],
            ["Over-broad consent", "Citizens may approve without understanding cross-agency scope.", "Purpose-specific receipts, preview, expiry, revocation; DPO."],
            ["Legacy API instability", "Retries can duplicate actions or create inconsistent states.", "Adapter idempotency, circuit breakers, reconciliation; integration owner."],
            ["Biased triage", "Reports or assistance may be deprioritized unfairly.", "Human rules, subgroup evaluation, appeal, no sole automated adverse decision."],
            ["Digital exclusion", "People without modern devices or literacy can be left behind.", "Assisted channels, low-bandwidth design, printable/SMS results; service owner."],
            ["Scope creep", "A super-agent can accumulate excessive data and authority.", "Risk tiers, architecture review, per-service approval, kill switch; steering group."],
            ["Vendor lock-in", "Models or workflow services may become hard to replace.", "Open contracts, model gateway, portable data, exit plan; enterprise architecture."],
        ],
        [0.21, 0.37, 0.42],
        font_size=7.45,
    )
    add_heading(doc, "Decision rights", 2)
    add_data_table(
        doc,
        ["Decision", "Accountable authority"],
        [
            ["Whether a service may be automated", "Source agency plus DICT governance and legal/privacy review"],
            ["What personal data may be processed", "Personal Information Controller, DPO, approved PIA/data-sharing terms"],
            ["Whether a transaction is complete", "Source agency system of record"],
            ["Whether a payment is settled", "eGovPay evidence plus agency ledger"],
            ["Whether an eReport is acknowledged/resolved", "Receiving agency or LGU"],
            ["Which model may be used", "Government AI/security governance under approved risk tier"],
        ],
        [0.42, 0.58],
        font_size=7.75,
    )
    add_callout(doc, "Human accountability", "The agent can recommend, prepare, and execute approved steps. It cannot become the legal decision-maker for adverse eligibility, enforcement, emergency prioritization, or issuance unless a specific law and agency policy authorize that automation.", tone="red", compact=True)

    # 15 Verification
    new_page(doc)
    page_label(doc, "14", "Verification and success metrics")
    add_heading(doc, "How technical judges - and production owners - should test it", 1)
    add_heading(doc, "Pre-production verification", 2)
    for text in [
        "Contract tests against every adapter sandbox, including version drift and malformed responses.",
        "Replay tests for timeouts after success, duplicated webhooks, stale tokens, partial fan-out, and restart recovery.",
        "Security testing: threat model, SAST, dependency scanning, API fuzzing, penetration test, secrets audit, prompt-injection red team.",
        "Privacy testing: data-flow inventory, minimization, consent withdrawal, deletion/retention, subject-access response, log redaction.",
        "AI evaluation: intent accuracy, extraction accuracy, tool-selection precision, unsafe-action rate, bilingual cases, refusal/handoff quality.",
        "Accessibility and resilience: keyboard/screen reader, 200% zoom, low bandwidth, device loss, offline draft, assisted channel.",
    ]:
        add_bullet(doc, text, bullet_id, size=8.75)
    add_heading(doc, "Pilot scorecard", 2)
    add_data_table(
        doc,
        ["Metric", "Target / guardrail"],
        [
            ["Correct service routing", ">=95% on adjudicated pilot set"],
            ["End-to-end completion", ">=70% for eligible, supported journeys; measure reasons for fallback"],
            ["Duplicate side effects", "0 confirmed duplicate bookings, reports, or charges"],
            ["Payment reconciliation", ">=99.95% automatically matched; all exceptions resolved within one business day"],
            ["Consent comprehension", ">=90% can identify agency, action, data, and fee in usability test"],
            ["Critical security/privacy incidents", "0; tested notification and containment procedure"],
            ["Accessibility", "WCAG 2.2 AA audit with no critical blockers"],
            ["Citizen time", "Measured reduction versus baseline, segmented by channel and accessibility need"],
        ],
        [0.41, 0.59],
        font_size=7.85,
    )
    add_callout(doc, "Go-live rule", "A visually impressive conversation is not a readiness signal. Go live only after system-of-record verification, failure recovery, privacy/security controls, agency UAT, support ownership, and measurable citizen benefit all pass.", tone="green")

    # 16 conclusion and references
    new_page(doc)
    page_label(doc, "15", "Conclusion and references")
    add_heading(doc, "Conclusion", 1)
    add_body(doc, "eGov Agent is technically feasible because it does not require government to start over. The necessary rails - a one-stop eGovPH channel, National ID authentication, data-exchange infrastructure, government payment capabilities, eReceipt, eReport, and agency systems - already exist in varying degrees. The missing product layer is safe orchestration: durable workflows, explicit consent, stable adapters, policy enforcement, authoritative read-back, and one citizen-facing trail.")
    add_body(doc, "The current prototype successfully communicates that future. Its next milestone is not a more powerful chatbot; it is a narrow, governed pilot with real agency owners, sandbox contracts, measurable outcomes, and production controls. If that pilot proves reliable, the adapter model can scale service by service without centralizing agency records or granting a model uncontrolled authority.")
    add_callout(doc, "Bottom line", "Keep the conversation simple for citizens and the controls rigorous underneath. The agent should make government feel unified while preserving where identity, money, records, and legal decisions actually belong.", tone="blue")
    add_heading(doc, "References", 2)
    refs = [
        ("[1] DICT eGov Data Exchange Platform (eGovDX), platform overview.", "https://core.e.gov.ph/"),
        ("[2] eGovPay, Payment Gateway API and Payment Links overview.", "https://egovpay.gov.ph/"),
        ("[3] Philippine Statistics Authority, rollout of Digital National ID and National ID eVerify, 10 June 2024.", "https://psa.gov.ph/content/psa-dict-roll-out-digital-national-id-authentication-services-national-id-everify-national"),
        ("[4] National ID eVerify, relying-party terms and onboarding context.", "https://everify.gov.ph/terms-of-use"),
        ("[5] National Privacy Commission, Republic Act No. 10173 - Data Privacy Act of 2012.", "https://privacy.gov.ph/data-privacy-act/"),
        ("[6] National Privacy Commission, Circular No. 2023-06 and official FAQ - Security of Personal Data.", "https://privacy.gov.ph/pips-and-pics/advisories-circulars/"),
        ("[7] Lawphil, Republic Act No. 11032 - Ease of Doing Business and Efficient Government Service Delivery Act of 2018.", "https://lawphil.net/statutes/repacts/ra2018/ra_11032_2018.html"),
        ("[8] DICT Citizen's Charter 2025, Integration in eGovPH Application, pp. 42-45.", "https://cms-cdn.e.gov.ph/DICT/pdf/DICT-Citizens-Charter-2025-1st-Edition.pdf"),
    ]
    for label, url in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.first_line_indent = Pt(-8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(label + " ")
        set_font(r, size=7.7, color=INK)
        add_hyperlink(p, url, url)
    add_heading(doc, "Document basis", 2)
    add_body(doc, f"Repository audit: eGov Agent prototype build {BUILD}, reviewed {DATE}. Public sources above were accessed {DATE}. Architecture, cost ranges, SLOs, and pilot targets are proposed engineering specifications or illustrative planning assumptions, not claims about current government production deployments.", size=8.0, color=MUTED)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()

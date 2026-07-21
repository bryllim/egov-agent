from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/bryllim/Developer/egov-agent")
OUT = ROOT / "output/docx/eGov-Agent-Technical-Whitepaper-One-Page.docx"
ASSET_DIR = ROOT / "tmp/whitepaper/assets_one_page"

# Design preset: compact_reference_guide.
# Named override: compact_reference_guide_a4_bw_one_page.
# The override is intentionally monochrome, A4, two-column, and print-dense.
FONT = "Arial"
BLACK = "000000"
GRAY = "555555"
LIGHT_GRAY = "E8E8E8"
WHITE = "FFFFFF"
PAGE_W_CM = 21.0
PAGE_H_CM = 29.7
MARGIN_CM = 1.15
HEADER_CM = 0.55
FOOTER_CM = 0.55
CONTENT_DXA = 10611
COLUMN_GAP_DXA = 300
COLUMN_DXA = 5155
TABLE_INDENT_DXA = 90
CELL_MARGINS_DXA = {"top": 50, "bottom": 50, "start": 90, "end": 90}
ARIAL = "/System/Library/Fonts/Supplemental/Arial.ttf"
ARIAL_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_font(run, *, size=None, bold=None, italic=None, color=BLACK) -> None:
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), FONT)
    r_pr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)


def image_font(size: int, *, bold=False):
    return ImageFont.truetype(ARIAL_BOLD if bold else ARIAL, size)


def centered_text(draw, box, text, font, *, fill="black") -> None:
    left, top, right, bottom = box
    bounds = draw.textbbox((0, 0), text, font=font)
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.text(((left + right - width) / 2, (top + bottom - height) / 2 - bounds[1]), text, font=font, fill=fill)


def arrow_down(draw, x: int, y1: int, y2: int, *, width=5) -> None:
    draw.line((x, y1, x, y2 - 14), fill="black", width=width)
    draw.polygon([(x, y2), (x - 11, y2 - 16), (x + 11, y2 - 16)], fill="black")


def arrow_right(draw, x1: int, y: int, x2: int, *, width=5) -> None:
    draw.line((x1, y, x2 - 14, y), fill="black", width=width)
    draw.polygon([(x2, y), (x2 - 16, y - 10), (x2 - 16, y + 10)], fill="black")


def make_architecture_illustration(path: Path) -> None:
    canvas = Image.new("L", (1180, 420), 255)
    draw = ImageDraw.Draw(canvas)
    title_font = image_font(29, bold=True)
    detail_font = image_font(22)
    boxes = [
        (35, 12, 1145, 86, "ANY CHANNEL", "Web  ·  Viber  ·  agency app  ·  API"),
        (35, 112, 1145, 198, "HEADLESS eGOV AGENT", "One core engine  ·  rules  ·  direct Gemini API"),
        (35, 224, 1145, 310, "GOVERNMENT SERVICES", "eVerify  ·  eGovPH  ·  eGovPay  ·  agency / LGU systems"),
        (35, 336, 1145, 410, "ARTIFACTS BACK IN CHAT", "Forms  ·  cards  ·  file previews  ·  receipts"),
    ]
    for index, (x1, y1, x2, y2, title, detail) in enumerate(boxes):
        fill = 0 if index in (1, 3) else 255
        ink = 255 if index == 1 else 0
        if index == 3:
            ink = 255
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=0, width=4)
        draw.text((x1 + 26, y1 + 13), title, font=title_font, fill=ink)
        bounds = draw.textbbox((0, 0), detail, font=detail_font)
        draw.text((x2 - 26 - (bounds[2] - bounds[0]), y1 + 20), detail, font=detail_font, fill=ink)
    arrow_down(draw, 590, 87, 110, width=4)
    arrow_down(draw, 590, 199, 222, width=4)
    arrow_down(draw, 590, 311, 334, width=4)
    canvas.save(path, optimize=True)


def make_feature_overview_illustration(path: Path) -> None:
    canvas = Image.new("L", (2200, 500), 255)
    draw = ImageDraw.Draw(canvas)
    cards = [
        ("UI", "GENERATIVE UI", "Builds useful cards, forms and receipts"),
        ("R", "HEURISTICS TIER", "Fast rules answer common requests first"),
        ("AI", "LOCAL + CLOUD AI", "Uses bigger AI only when needed"),
        ("<> ", "HEADLESS", "Works in chat, Viber, apps and APIs"),
        ("A", "AGENTIC", "Plans, asks, acts and returns proof"),
        ("VM", "PRIVATE TASK MACHINE", "Complex jobs run in a temporary sandbox"),
    ]
    card_w, card_h = 690, 205
    x_positions = [35, 755, 1475]
    y_positions = [25, 270]
    for index, (icon, title, detail) in enumerate(cards):
        col, row = index % 3, index // 3
        x1, y1 = x_positions[col], y_positions[row]
        x2, y2 = x1 + card_w, y1 + card_h
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=255, outline=0, width=4)
        draw.ellipse((x1 + 24, y1 + 35, x1 + 138, y1 + 149), fill=0)
        centered_text(draw, (x1 + 24, y1 + 35, x1 + 138, y1 + 149), icon.strip(), image_font(34, bold=True), fill="white")
        draw.text((x1 + 165, y1 + 43), title, font=image_font(31, bold=True), fill=0)
        draw.text((x1 + 165, y1 + 96), detail, font=image_font(23), fill="#555555")
    canvas.save(path, optimize=True)


def make_tier_illustration(path: Path) -> None:
    canvas = Image.new("L", (1180, 315), 255)
    draw = ImageDraw.Draw(canvas)
    rows = [
        ("1", "FAST RULES", "Known requests", "78%  ·  <10 ms  ·  ~P0"),
        ("2", "LOCAL AI", "Taglish + data extraction", "17%  ·  ~120 ms  ·  ~P0.02"),
        ("3", "CLOUD AI", "Hard planning; private fields removed", "5%  ·  1-3 s  ·  ~P1.50"),
    ]
    for index, (number, title, detail, metric) in enumerate(rows):
        y1 = 8 + index * 102
        y2 = y1 + 88
        fill = 0 if index == 0 else 255
        ink = 255 if index == 0 else 0
        draw.rounded_rectangle((24, y1, 1156, y2), radius=17, fill=fill, outline=0, width=4)
        draw.ellipse((43, y1 + 11, 109, y1 + 77), fill=255 if index == 0 else 0, outline=0)
        centered_text(draw, (43, y1 + 11, 109, y1 + 77), number, image_font(27, bold=True), fill="black" if index == 0 else "white")
        draw.text((135, y1 + 14), title, font=image_font(27, bold=True), fill=ink)
        draw.text((135, y1 + 49), detail, font=image_font(21), fill=ink if index == 0 else "#555555")
        bounds = draw.textbbox((0, 0), metric, font=image_font(22, bold=True))
        draw.text((1128 - (bounds[2] - bounds[0]), y1 + 30), metric, font=image_font(22, bold=True), fill=ink)
    canvas.save(path, optimize=True)


def make_guardrail_illustration(path: Path) -> None:
    canvas = Image.new("L", (1180, 205), 255)
    draw = ImageDraw.Draw(canvas)
    label_font = image_font(25, bold=True)
    role_font = image_font(20)
    centers = [100, 420, 740, 1060]
    labels = [("PLAN", "agent"), ("APPROVE", "person + policy"), ("DO", "agency"), ("CONFIRM", "official source")]
    for index, (x, (label, role)) in enumerate(zip(centers, labels), start=1):
        if index < 4:
            arrow_right(draw, x + 58, 74, centers[index] - 58, width=4)
        draw.ellipse((x - 52, 22, x + 52, 126), fill=0)
        centered_text(draw, (x - 52, 22, x + 52, 126), str(index), image_font(40, bold=True), fill="white")
        centered_text(draw, (x - 130, 132, x + 130, 166), label, label_font)
        centered_text(draw, (x - 140, 167, x + 140, 200), role, role_font, fill="#555555")
    canvas.save(path, optimize=True)


def make_payment_illustration(path: Path) -> None:
    canvas = Image.new("L", (1180, 235), 255)
    draw = ImageDraw.Draw(canvas)
    label_font = image_font(22, bold=True)
    centers = [90, 340, 590, 840, 1090]
    labels = ["ORDER", "eGOVPAY", "VERIFY", "POST ONCE", "eRECEIPT"]
    for index, (x, label) in enumerate(zip(centers, labels), start=1):
        if index < 5:
            arrow_right(draw, x + 45, 88, centers[index] - 45, width=4)
        fill = 0 if index in (1, 3, 5) else 255
        ink = 255 if fill == 0 else 0
        draw.ellipse((x - 45, 43, x + 45, 133), fill=fill, outline=0, width=4)
        centered_text(draw, (x - 45, 43, x + 45, 133), str(index), image_font(34, bold=True), fill="white" if ink == 255 else "black")
        centered_text(draw, (x - 105, 150, x + 105, 188), label, label_font)
    draw.rounded_rectangle((190, 202, 990, 232), radius=13, fill=0)
    centered_text(draw, (190, 202, 990, 232), "HOSTED CHECKOUT  ·  SERVER VERIFICATION  ·  AUTHORITATIVE PROOF", image_font(17, bold=True), fill="white")
    canvas.save(path, optimize=True)


def generate_illustrations() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "architecture": ASSET_DIR / "architecture.png",
        "features": ASSET_DIR / "features.png",
        "tiers": ASSET_DIR / "tiers.png",
        "guardrails": ASSET_DIR / "guardrails.png",
        "payment": ASSET_DIR / "payment.png",
    }
    make_architecture_illustration(paths["architecture"])
    make_feature_overview_illustration(paths["features"])
    make_tier_illustration(paths["tiers"])
    make_guardrail_illustration(paths["guardrails"])
    make_payment_illustration(paths["payment"])
    return paths


def paragraph_border(paragraph, *, side="bottom", size=8, space=4, color=BLACK) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def shade_paragraph(paragraph, fill=BLACK) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, *, indent=TABLE_INDENT_DXA) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, *, color=BLACK, size=4, inside=True) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "start", "bottom", "end"]
    if inside:
        names += ["insideH", "insideV"]
    for name in names:
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)
        borders.append(edge)


def configure_numbering(doc: Document) -> tuple[int, int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    base_abs = max(existing_abs or [0]) + 1
    base_num = max(existing_num or [0]) + 1

    def make_abstract(abs_id: int, fmt: str, text: str, left: int, hanging: int) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        for tag, value in (("w:start", "1"), ("w:numFmt", fmt), ("w:lvlText", text), ("w:lvlJc", "left")):
            node = OxmlElement(tag)
            node.set(qn("w:val"), value)
            lvl.append(node)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "30")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def make_num(num_id: int, abs_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abs_id))
        num.append(ref)
        numbering.append(num)

    make_abstract(base_abs, "bullet", "•", 270, 135)
    make_num(base_num, base_abs)
    make_abstract(base_abs + 1, "decimal", "%1.", 300, 170)
    make_num(base_num + 1, base_abs + 1)
    # A distinct abstract list guarantees the payment sequence restarts at 1
    # in both Word and LibreOffice.
    make_abstract(base_abs + 2, "decimal", "%1.", 300, 170)
    make_num(base_num + 2, base_abs + 2)
    return base_num, base_num + 1, base_num + 2


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def set_section_geometry(section) -> None:
    section.page_width = Cm(PAGE_W_CM)
    section.page_height = Cm(PAGE_H_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(HEADER_CM)
    section.footer_distance = Cm(FOOTER_CM)


def set_two_columns(section) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(COLUMN_GAP_DXA))
    cols.set(qn("w:equalWidth"), "1")


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(8.0)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.0)
    normal.paragraph_format.line_spacing = 1.03

    h1 = styles["Heading 1"]
    h1.font.name = FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h1.font.size = Pt(10.2)
    h1.font.bold = True
    h1.font.color.rgb = rgb(BLACK)
    h1.paragraph_format.space_before = Pt(5.0)
    h1.paragraph_format.space_after = Pt(2.2)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2.font.size = Pt(8.8)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BLACK)
    h2.paragraph_format.space_before = Pt(3.5)
    h2.paragraph_format.space_after = Pt(1.5)
    h2.paragraph_format.keep_with_next = True


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(1.5)
    r = kicker.add_run("TECHNICAL BRIEF  |  22 JULY 2026")
    set_font(r, size=7.3, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1.2)
    r = title.add_run("eGov Agent")
    set_font(r, size=20.5, bold=True)
    r = title.add_run("  |  Government services through one conversation")
    set_font(r, size=12.3, bold=False, color=GRAY)

    thesis = doc.add_paragraph()
    thesis.paragraph_format.space_after = Pt(4.5)
    thesis.paragraph_format.line_spacing = 1.0
    r = thesis.add_run(
        "eGovPH is already great. eGov Agent makes it easier: tell it what you need, and it uses the government connections that already exist, asks before acting, and gives you clear proof when the job is done."
    )
    set_font(r, size=8.7, bold=True)
    paragraph_border(thesis, size=10, space=4)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 1")
    paragraph_border(p, size=5, space=2)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, after=2.0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, size=8.0, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, size=8.0)
    else:
        r = p.add_run(text)
        set_font(r, size=8.0)


def add_bullet(doc: Document, text: str, bullet_id: int, *, size=7.8, after=1.3) -> None:
    p = doc.add_paragraph()
    apply_num(p, bullet_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_font(r, size=size)


def add_number(doc: Document, text: str, number_id: int, *, size=7.8, after=1.3) -> None:
    p = doc.add_paragraph()
    apply_num(p, number_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_font(r, size=size)


def add_status_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "STATUS"
    table.rows[0].cells[1].text = "WHAT IT MEANS"
    rows = [
        ("Working now", "Chat, live artifacts, file previews, uploads, memory, vault and printable results"),
        ("Simulated", "eVerify, agency records, appointments, eGovPay/eReceipt and eReport routing"),
        ("Still needed", "Live agency APIs, backend, database, job queue, credentials and payment callbacks"),
    ]
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1450, COLUMN_DXA - 1450])
    set_table_borders(table, size=4)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, BLACK)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=6.9 if r_idx == 0 else 7.2, bold=(r_idx == 0 or c_idx == 0), color=WHITE if r_idx == 0 else BLACK)


def add_architecture(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.0
    parts = [
        ("Channels", True), ("  →  ", False),
        ("API gateway", True), ("  →  ", False),
        ("Agent orchestrator", True), ("  →  ", False),
        ("Policy + consent", True), ("  →  ", False),
        ("Versioned adapters", True), ("  →  ", False),
        ("eVerify / eGovDX / eGovPay / agency-LGU systems", True),
    ]
    for text, bold in parts:
        r = p.add_run(text)
        set_font(r, size=7.6, bold=bold)


def add_inline_figure(doc: Document, path: Path, *, width_in: float, alt_text: str, after=2.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width_in))
    shape._inline.docPr.set("descr", alt_text)


def add_callout(doc: Document, title: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(4)
    p.paragraph_format.right_indent = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    shade_paragraph(p, BLACK)
    r = p.add_run(title.upper() + "  ")
    set_font(r, size=7.6, bold=True, color=WHITE)
    r = p.add_run(text)
    set_font(r, size=7.6, bold=False, color=WHITE)


def add_pilot_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "WINDOW"
    table.rows[0].cells[1].text = "DELIVERABLE"
    rows = [
        ("0-6 weeks", "Rules, privacy review, service maps and connector plans"),
        ("6-12 weeks", "Read-only data checks, identity, consent and test connections"),
        ("12-24 weeks", "Pilot payments, eReceipts, reports and human help"),
        ("6-12 months", "2 national agencies + 1 LGU; security and wider rollout"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1300, COLUMN_DXA - 1300])
    set_table_borders(table, size=4)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, BLACK)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, size=6.8 if r_idx == 0 else 7.0, bold=(r_idx == 0 or c_idx == 0), color=WHITE if r_idx == 0 else BLACK)


def set_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("eGov Agent  |  One-page technical whitepaper  |  Feasibility brief, not a procurement quote")
    set_font(r, size=6.3, color=GRAY)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    illustrations = generate_illustrations()
    doc = Document()
    setup_styles(doc)
    first = doc.sections[0]
    set_section_geometry(first)
    set_footer(first)
    bullet_id, number_id, payment_number_id = configure_numbering(doc)

    props = doc.core_properties
    props.title = "eGov Agent - One-Page Technical Whitepaper"
    props.subject = "Feasible agentic orchestration layer for eGovPH"
    props.author = "eGov Agent Team"
    props.keywords = "eGovPH, agentic government, eGovPay, eReceipt, eReport, interoperability"

    add_title_block(doc)
    add_inline_figure(
        doc,
        illustrations["features"],
        width_in=7.03,
        alt_text="Six main features: generative UI, heuristics tier, local and cloud AI, headless channels, agentic action, and private task machines.",
        after=2.2,
    )

    body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_section_geometry(body_section)
    set_two_columns(body_section)
    set_footer(body_section)

    add_heading(doc, "1  THE SIMPLE FLOW")
    add_inline_figure(
        doc,
        illustrations["architecture"],
        width_in=3.48,
        alt_text="A request can come from any channel, run through the headless eGov Agent, use government services, and return useful artifacts in the chat.",
        after=1.5,
    )
    add_bullet(doc, "Headless means the core agent is not tied to this website. The same logic can work in web chat, Viber, an agency app or an API.", bullet_id)
    add_bullet(doc, "Official records stay with the agency. The agent connects services and brings the result back to one conversation.", bullet_id)

    add_heading(doc, "2  THE HEURISTICS TIER SAVES MONEY")
    add_inline_figure(
        doc,
        illustrations["tiers"],
        width_in=3.48,
        alt_text="Projected request routing: 78 percent fast rules, 17 percent structured workflows, and 5 percent complex AI.",
        after=1.2,
    )
    add_body(doc, "The router starts with the cheapest safe option and moves up only when confidence is low. The percentages and prices are planning assumptions for the prototype.", after=1.2)
    add_bullet(doc, "Heuristics tier: simple rules answer known requests and status checks almost instantly.", bullet_id)
    add_bullet(doc, "Local AI: understands Taglish and pulls fields from text while private data stays on government infrastructure.", bullet_id)
    add_bullet(doc, "Cloud AI: used only for hard planning; names, IDs and other private fields are removed first.", bullet_id)

    add_heading(doc, "3  PERSONAL, BUT USER-CONTROLLED")
    add_bullet(doc, "Memory keeps approved facts and preferences so people do not repeat the same details. Every fact needs a source and an edit/delete option.", bullet_id)
    add_bullet(doc, "The Vault stores encrypted documents and shows friendly file-stamp previews when a form needs an attachment.", bullet_id)
    add_bullet(doc, "Long tasks can run in a temporary private machine. It is created for one job, returns the files, then is wiped.", bullet_id)

    add_heading(doc, "4  WHAT THE PROTOTYPE PROVES")
    add_status_table(doc)

    # Explicit column break keeps the one-page composition stable across renderers.
    column_break = doc.add_paragraph()
    column_break.paragraph_format.space_after = Pt(0)
    column_break.add_run().add_break(WD_BREAK.COLUMN)

    add_heading(doc, "5  GENERATIVE UI (ARTIFACTS)")
    add_body(doc, "The agent does not answer with text alone. It builds the right interface inside the conversation. We call each generated result an artifact.", after=1.2)
    add_bullet(doc, "Review cards and pre-filled forms before anything is submitted.", bullet_id)
    add_bullet(doc, "File-stamp previews for IDs, certificates, photos and uploaded documents.", bullet_id)
    add_bullet(doc, "eReport triage, responder status, appointment passes, payment cards and print-ready eReceipts.", bullet_id)
    add_callout(doc, "Why it matters", "Less reading, fewer screens and a clear result people can save or print.")

    add_heading(doc, "6  AGENTIC, WITH CLEAR GUARDRAILS")
    add_inline_figure(
        doc,
        illustrations["guardrails"],
        width_in=3.48,
        alt_text="Safe action sequence: plan, authorize, execute, then confirm from the official agency system.",
        after=1.0,
    )
    add_bullet(doc, "The agent shows the plan, data, fees and agencies first. The person approves before a booking, report or payment.", bullet_id)
    add_bullet(doc, "A screen that says success is not enough. The backend checks the official agency system and keeps an audit trail.", bullet_id)

    add_heading(doc, "7  eGOVPAY TO VERIFIED eRECEIPT")
    add_inline_figure(
        doc,
        illustrations["payment"],
        width_in=3.48,
        alt_text="Verified payment path: create order, use hosted eGovPay, verify server-side, post once, and issue the authoritative eReceipt.",
        after=1.2,
    )
    add_bullet(doc, "The agent never stores card or wallet details. The backend checks the order, amount and merchant, and blocks duplicate events.", bullet_id, size=7.6)
    add_bullet(doc, "Only the agency or eReceipt service creates the official receipt after the payment is posted and checked.", bullet_id, size=7.6)

    add_heading(doc, "8  A PRACTICAL PILOT")
    add_pilot_table(doc)
    add_body(doc, "Rough 9-12 month planning range: PHP 18.0M-33.5M for two national agencies, one LGU and up to 50,000 monthly users. This is not a vendor quote.", bold_lead="Rough 9-12 month planning range:", after=1.2)
    add_bullet(doc, "Targets: at least 95% correct routing, no duplicate high-impact actions, at least 99.95% payment matching and WCAG 2.2 AA.", bullet_id, size=7.4)
    add_bullet(doc, "Before launch: Privacy Impact Assessment, security testing, agency approval and a human path for uncertain or high-risk cases.", bullet_id, size=7.4)

    add_heading(doc, "OFFICIAL BASIS")
    refs = doc.add_paragraph()
    refs.paragraph_format.space_after = Pt(1)
    refs.paragraph_format.line_spacing = 1.0
    ref_text = (
        "eGovPH Core (core.e.gov.ph) · eGovPay (egovpay.gov.ph) · PSA National ID eVerify (psa.gov.ph; everify.gov.ph) · "
        "Data Privacy Act and NPC Circular 2023-06 (privacy.gov.ph) · RA 11032 (lawphil.net)"
    )
    r = refs.add_run(ref_text)
    set_font(r, size=6.6, color=GRAY)
    add_callout(doc, "Recommendation", "Start with a small sandbox pilot. Keep agencies in control, ask before every important action, and confirm every result from the official source.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
